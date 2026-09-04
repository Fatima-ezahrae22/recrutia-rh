# -*- coding: utf-8 -*-
"""
generer_rapport_docx.py
Génère le document Rapport_de_Stage_Agent_IA_RH_ArtiWeb.docx
"""
import sys, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Color Palette
COLOR_PRIMARY = RGBColor(0, 51, 102)     # Deep Navy Blue
COLOR_SECONDARY = RGBColor(0, 102, 204)  # Royal Blue
COLOR_TEXT = RGBColor(51, 51, 51)        # Dark Gray
COLOR_MUTED = RGBColor(102, 102, 102)

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    p.paragraph_format.space_after = Pt(6)
    return p

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = COLOR_SECONDARY
    p.paragraph_format.space_after = Pt(24)
    return p

def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLOR_SECONDARY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_p(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_TEXT
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_TEXT
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_TEXT
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_TEXT
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(40, 40, 40)
    return p

# --- TITLE ---
add_title("RAPPORT DE STAGE ET DE RÉALISATION TECHNIQUE")
add_subtitle("Développement de la Couche d'Ingestion Intelligente de CV (PDF / DOCX / OCR)\nProjet Agent IA RH (RecrutIA) — Entreprise d'Accueil : ArtiWeb")

# --- SECTION 1 ---
add_h1("I. PRÉSENTATION DE L'ENTREPRISE D'ACCUEIL : ARTIWEB")

add_h2("1.1 Présentation Générale")
add_bullet(" ARTIWEB (ARTI web)", "Nom de l'entreprise :")
add_bullet(" https://artiweb.ma/", "Site Web Officiel :")
add_bullet(" Fès, Maroc", "Siège Social :")
add_bullet(" Partenaire Google officiel (Google Partner)", "Partenariats :")
add_bullet(" Lundi – Vendredi (9:00 – 21:00)", "Horaires :")
add_bullet(" (+212) 664-017447 | contact@artiweb.ma", "Contact :")

add_h2("1.2 Secteurs d'Activité et Expertises")
add_p("ArtiWeb est une agence de communication et de marketing digital basée à Fès. Elle accompagne les entreprises dans la transformation digitale et l'optimisation de leur présence en ligne à travers plusieurs pôles d'expertise :")
add_bullet(" Conduite de campagnes de communication, branding et visibilité digitale.", "Stratégie Digitale & Marketing Numérique :")
add_bullet(" Conception de sites web sur-mesure, plateformes e-commerce et applications métiers.", "Développement Web & Applications :")
add_bullet(" Positionnement SEO sur moteurs de recherche, gestion de campagnes Google Ads et Social Ads.", "Référencement & Acquisition :")
add_bullet(" Création visuelle, contenu multimédia et optimisation des canaux de vente e-commerce.", "Studio Marketing & Marketplace :")

add_p("Dans le cadre de son développement et de la modernisation des outils RH internes et clients, ArtiWeb a lancé le projet Agent IA RH (RecrutIA).")

# --- SECTION 2 ---
add_h1("II. FICHE PROJET : AGENT IA RH (RECRUTIA)")

add_h2("2.1 Intitulé du Projet")
add_p("Agent IA RH — Système intelligent d'automatisation du processus de recrutement (Nom court : RecrutIA).")

add_h2("2.2 Problématique")
add_p("Le recrutement traditionnel repose sur le tri manuel de chaque CV (lecture individuelle, comparaison aux exigences du poste, décision au cas par cas). Cette méthode devient un goulot d'étranglement avec un volume élevé de candidatures : elle est chronophage, sujette aux erreurs humaines et biais, et manque de traçabilité.")
add_p("Problématique centrale : Comment automatiser l'analyse et le tri des candidatures pour réduire la charge de travail des RH, tout en garantissant que les décisions restent fiables, explicables et sous contrôle humain sur les cas sensibles ?")

add_h2("2.3 Architecture en 3 Couches")
add_bullet(" Extrait et structure automatiquement les données des CV bruts (Texte, OCR, Compétences, Expérience, Formation) sans dépendre d'un LLM.", "Couche 1 - Ingestion (Tâche réalisée) :")
add_bullet(" Compare sémantiquement chaque profil à l'offre d'emploi, calcule un score de correspondance et génère une justification argumentée.", "Couche 2 - IA Agentique :")
add_bullet(" Le recruteur garde la décision finale (validation, correction ou rejet) sur chaque préconisation de l'agent (principe Human-in-the-loop).", "Couche 3 - Supervision Humaine :")

# --- SECTION 3 ---
add_h1("III. DESCRIPTION DÉTAILLÉE DE LA PREMIÈRE TÂCHE : COUCHE D'INGESTION")

add_h2("3.1 Objectif de la Tâche")
add_p("Développer un module Python autonome capable de transformer un CV brut (fichier PDF textuel, PDF scanné/image ou DOCX Word) en un dictionnaire structuré JSON exploitable par la couche IA.")

add_h2("3.2 Pourquoi c'est la première tâche")
add_bullet("Elle ne dépend d'aucune autre couche (pas de LLM nécessaire) -> développée et testée en totale autonomie.", "Autonomie :")
add_bullet("Elle produit le contrat de données (JSON) indispensable pour les couches IA et Backend.", "Contrat de données :")
add_bullet("Elle traite le risque majeur de qualité d'extraction sur des CV mal formatés ou scannés dès le début du projet.", "Réduction des risques :")

add_h2("3.3 Architecture des Fichiers du Module (ingestion/)")
add_code("""agent_ai/
├── ingestion/
│   ├── __init__.py        # Export de la fonction principale traiter_cv()
│   ├── extractor.py       # Extraction texte PDF/DOCX + Fallback OCR Tesseract
│   ├── cleaner.py         # Nettoyage, normalisation Unicode et suppression du bruit
│   ├── parser.py          # Extraction NLP des compétences, expérience et formation
│   └── traiter_cv.py      # Orchestrateur central du pipeline
├── data/
│   └── competences_ref.json # Référentiel de ~80 compétences IT & Soft Skills
├── mes_cv/                # Dossier de dépôt des CV réels
├── tests/
│   └── test_ingestion.py  # Suite de 21 tests unitaires et d'intégration
├── run.py                 # Démonstration du pipeline
├── analyser_mon_cv.py     # Script interactif d'analyse de CV réels
└── requirements.txt       # Dépendances Python""")

add_h2("3.4 Technologies & Libraries Utilisées")
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Technologie / Library"
hdr_cells[1].text = "Rôle dans le projet"
set_cell_background(hdr_cells[0], "003366")
set_cell_background(hdr_cells[1], "003366")
for cell in hdr_cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

data_tech = [
    ("Python 3.12", "Langage principal de développement backend."),
    ("pdfplumber", "Extraction haute précision du texte natif et de la structure des PDF."),
    ("PyMuPDF (fitz)", "Extraction ultra-rapide et conversion des pages PDF en images HD (300 DPI) pour l'OCR."),
    ("python-docx", "Extraction des paragraphes et des cellules de tableaux dans les fichiers Word (.docx)."),
    ("pytesseract + Tesseract OCR", "Moteur OCR de Google pour lire le texte dans les CV scannés/images."),
    ("unicodedata & re", "Normalisation Unicode NFC, suppression du bruit et expressions régulières (Regex)."),
    ("pytest & unittest", "Frameworks de tests automatisés (21 tests).")
]

for tech, role in data_tech:
    row_cells = table.add_row().cells
    row_cells[0].text = tech
    row_cells[1].text = role

add_h2("3.5 Le Pipeline en 4 Étapes-Clés")
add_bullet("Tentative native via pdfplumber / python-docx -> Fallback PyMuPDF -> Fallback OCR Tesseract si le texte fait < 80 caractères.", "1. Extraction en cascade (extractor.py) :")
add_bullet("Normalisation Unicode NFC, fusion des mots coupés (ex: 'Déve-\\nloppeur'), correction des ligatures PDF ('ﬁ' -> 'fi'), suppression des puces (•, ★) et en-têtes.", "2. Nettoyage & Normalisation (cleaner.py) :")
add_bullet("Extraction des compétences par matching avec competences_ref.json, calcul de l'expérience via mentions explicites ou fusion de dates, détection de la formation par hiérarchie.", "3. Parsing NLP (parser.py) :")
add_bullet("Orchestration et mesure du temps de traitement. Génération de l'objet JSON structuré.", "4. Assemblage JSON (traiter_cv.py) :")

# --- SECTION 4 ---
add_h1("IV. EXPLICATION COMPLÈTE DES TESTS AUTOMATISÉS")

add_p("Une suite de 21 tests automatisés a été développée dans tests/test_ingestion.py pour valider le code avant déploiement.")

add_h2("4.1 Structure des Tests (21 Tests au total)")
add_bullet("Vérifient isolément que le nettoyeur élimine les puces, corrige les ligatures PDF et gère les espaces.", "TestCleaner (5 tests unitaires) :")
add_bullet("Vérifient l'extraction des compétences, le calcul des durées d'expérience et la priorité des diplômes (Doctorat > Master > Bac).", "TestParser (9 tests unitaires) :")
add_bullet("Génèrent de vrais fichiers PDF et DOCX temporaires et font passer tout le pipeline de A à Z.", "TestIntegration (7 tests d'intégration) :")

add_h2("4.2 Résultat de la Suite de Tests")
add_code("""============================= test session starts =============================
platform win32 -- Python 3.12.9, pytest-9.1.1, pluggy-1.6.0
collected 21 items

tests/test_ingestion.py::TestCleaner::test_espaces_multiples PASSED
tests/test_ingestion.py::TestCleaner::test_ligatures_pdf PASSED
tests/test_ingestion.py::TestCleaner::test_normalisation_unicode PASSED
tests/test_ingestion.py::TestCleaner::test_suppression_caracteres_parasites PASSED
tests/test_ingestion.py::TestCleaner::test_texte_vide PASSED
tests/test_ingestion.py::TestParser::test_experience_calcul_dates PASSED
tests/test_ingestion.py::TestParser::test_experience_mention_explicite PASSED
tests/test_ingestion.py::TestParser::test_experience_zero_si_non_trouve PASSED
tests/test_ingestion.py::TestParser::test_extraction_competences_basique PASSED
tests/test_ingestion.py::TestParser::test_extraction_competences_insensible_casse PASSED
tests/test_ingestion.py::TestParser::test_formation_bts PASSED
tests/test_ingestion.py::TestParser::test_formation_ingenieur PASSED
tests/test_ingestion.py::TestParser::test_formation_master PASSED
tests/test_ingestion.py::TestParser::test_formation_non_specifiee PASSED
tests/test_ingestion.py::TestParser::test_formation_priorite_plus_haut_niveau PASSED
tests/test_ingestion.py::TestIntegrationTraiterCV::test_affichage_json_final PASSED
tests/test_ingestion.py::TestIntegrationTraiterCV::test_cv_docx PASSED
tests/test_ingestion.py::TestIntegrationTraiterCV::test_cv_pdf_bien_formate PASSED
tests/test_ingestion.py::TestIntegrationTraiterCV::test_cv_pdf_mal_formate PASSED
tests/test_ingestion.py::TestIntegrationTraiterCV::test_fichier_inexistant PASSED
tests/test_ingestion.py::TestIntegrationTraiterCV::test_format_non_supporte PASSED

============================== 21 passed in 0.91s ==============================""")

# --- SECTION 5 ---
add_h1("V. ANCIENNE & NOUVELLE ANIMALES : TEST SUR UN CV RÉEL")

add_h2("5.1 Analyse du CV Réel (MEKKI_FATIMA-EZAHRAE_AI_Backend.pdf)")
add_p("Le script analyser_mon_cv.py a été exécuté sur un vrai CV réel. Voici les résultats obtenus :")

add_bullet(" MEKKI_FATIMA-EZAHRAE_AI_Backend.pdf", "Fichier analysé :")
add_bullet(" SUCCÈS", "Statut :")
add_bullet(" Master / Ingénieur (Détecté 'Ingénieure en Génie Informatique')", "Formation :")
add_bullet(" 2 an(s) (Calculé automatiquement depuis les périodes)", "Expérience :")
add_bullet(" 0.782 seconde", "Durée de traitement :")
add_bullet(" 33 compétences détectées (Python, Django, Flask, TensorFlow, React, Azure, Java, MySQL, Node.js, Scikit-learn, etc.)", "Compétences :")

add_h2("5.2 Extrait du JSON généré (mes_cv/MEKKI_FATIMA-EZAHRAE_AI_Backend_resultat.json)")
add_code("""{
    "fichier": "MEKKI_FATIMA-EZAHRAE_AI_Backend.pdf",
    "statut": "succès",
    "formation": "Master / Ingénieur",
    "experience_annees": 2,
    "nb_competences": 33,
    "competences": [
        "Azure", "Bootstrap", "CSS", "Cloud", "Computer Vision", "DevOps",
        "Django", "Express.js", "Flask", "Git", "HTML", "Java", "JavaScript",
        "Laravel", "MySQL", "NLP", "Node.js", "PHP", "Python", "React",
        "REST API", "Scikit-learn", "Spring Boot", "TensorFlow"
    ],
    "message": "CV traité avec succès (33 compétence(s) détectée(s)).",
    "duree_traitement": "0.782 s"
}""")

# --- SECTION 6 ---
add_h1("VI. CONCLUSION ET PROCHAINES ÉTAPES")
add_p("Le module d'ingestion développé pour ArtiWeb est 100% fonctionnel, autonome, rapide (< 1s) et validé par 21 tests automatisés ainsi que sur un vrai CV réels. Il fournit le contrat de données JSON nécessaire à la couche d'Intelligence Artificielle Agentique (Couche 2).")

# Save document
output_docx = "Rapport_de_Stage_Agent_IA_RH_ArtiWeb.docx"
doc.save(output_docx)
print(f"Document DOCX généré avec succès : {os.path.abspath(output_docx)}")
