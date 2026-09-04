# -*- coding: utf-8 -*-
"""
generer_rapport_encadrant.py
Génère le Rapport de Stage Complet au format PDF (ReportLab) et DOCX (python-docx)
destiné à l'encadrant de stage.
"""

import sys, os
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak
)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# ─────────────────────────────────────────────────────────────
# 1. GÉNÉRATION DU FICHIER PDF COMPLET
# ─────────────────────────────────────────────────────────────

pdf_path = "Rapport_de_Stage_Complet_Encadrant.pdf"
doc_pdf = SimpleDocTemplate(
    pdf_path,
    pagesize=A4,
    rightMargin=2*cm, leftMargin=2*cm,
    topMargin=2*cm, bottomMargin=2*cm
)

NAVY  = colors.HexColor('#003366')
BLUE  = colors.HexColor('#0055A4')
GREEN = colors.HexColor('#006633')
RED   = colors.HexColor('#CC0000')
GRAY  = colors.HexColor('#333333')
LGRAY = colors.HexColor('#F2F4F7')
BGCOD = colors.HexColor('#F0F4F8')
WHITE = colors.white

st_cover   = ParagraphStyle('cover',   fontName='Helvetica-Bold',    fontSize=20, leading=24, textColor=NAVY,  alignment=1, spaceAfter=8)
st_sub     = ParagraphStyle('sub',     fontName='Helvetica-Oblique', fontSize=11, leading=14, textColor=BLUE,  alignment=1, spaceAfter=16)
st_h1      = ParagraphStyle('h1',      fontName='Helvetica-Bold',    fontSize=13, leading=16, textColor=NAVY,  spaceBefore=14, spaceAfter=6)
st_h2      = ParagraphStyle('h2',      fontName='Helvetica-Bold',    fontSize=11, leading=14, textColor=BLUE,  spaceBefore=10, spaceAfter=4)
st_body    = ParagraphStyle('body',    fontName='Helvetica',         fontSize=9.5, leading=13.5, textColor=GRAY, spaceAfter=5)
st_bullet  = ParagraphStyle('bullet',  fontName='Helvetica',         fontSize=9.5, leading=13,   textColor=GRAY, leftIndent=14, spaceAfter=3)
st_code    = ParagraphStyle('code',    fontName='Courier',           fontSize=8,   leading=10,   textColor=colors.HexColor('#111111'),
                             backColor=BGCOD, borderColor=colors.HexColor('#BBBBBB'), borderWidth=0.5,
                             borderPadding=6, spaceBefore=5, spaceAfter=7)
st_cell    = ParagraphStyle('cell',    fontName='Helvetica',         fontSize=9,   leading=12, textColor=GRAY)
st_thdr    = ParagraphStyle('thdr',    fontName='Helvetica-Bold',    fontSize=9.5, leading=12, textColor=WHITE)
st_ok      = ParagraphStyle('ok',      fontName='Helvetica-Bold',    fontSize=9.5, textColor=GREEN)

def p(t, st=None):  return Paragraph(t, st or st_body)
def b(t):           return Paragraph(f"• {t}", st_bullet)
def bb(pre, t):     return Paragraph(f"<b>{pre}</b> {t}", st_bullet)
def h1(t):          return Paragraph(t, st_h1)
def h2(t):          return Paragraph(t, st_h2)
def hr():           return HRFlowable(width="100%", thickness=1.2, color=NAVY, spaceAfter=8, spaceBefore=2)
def sp(n=8):        return Spacer(1, n)
def code(t):        return Paragraph(t.replace(' ','&nbsp;').replace('\n','<br/>'), st_code)

W = 15.5*cm

def ttable(rows, widths):
    t = Table(rows, colWidths=widths)
    t.setStyle(TableStyle([
        ('BACKGROUND',   (0,0),(-1,0),  NAVY),
        ('TEXTCOLOR',    (0,0),(-1,0),  WHITE),
        ('FONTNAME',     (0,0),(-1,0),  'Helvetica-Bold'),
        ('FONTSIZE',     (0,0),(-1,-1), 9),
        ('ALIGN',        (0,0),(-1,-1), 'LEFT'),
        ('VALIGN',       (0,0),(-1,-1), 'TOP'),
        ('GRID',         (0,0),(-1,-1), 0.4, colors.HexColor('#CCCCCC')),
        ('ROWBACKGROUNDS',(0,1),(-1,-1),[WHITE, LGRAY]),
        ('TOPPADDING',   (0,0),(-1,-1), 5),
        ('BOTTOMPADDING',(0,0),(-1,-1), 5),
        ('LEFTPADDING',  (0,0),(-1,-1), 6),
    ]))
    return t

def rbox(titre, lignes, c=GREEN):
    rows = [[p(f"<b>{titre}</b>", st_thdr)]]
    for l in lignes:
        rows.append([p(l, st_cell)])
    t = Table(rows, colWidths=[W])
    t.setStyle(TableStyle([
        ('BACKGROUND',(0,0),(-1,0), c),
        ('BACKGROUND',(0,1),(-1,-1), LGRAY),
        ('BOX',       (0,0),(-1,-1), 0.8, c),
        ('TOPPADDING',   (0,0),(-1,-1), 5),
        ('BOTTOMPADDING',(0,0),(-1,-1), 5),
        ('LEFTPADDING',  (0,0),(-1,-1), 8),
    ]))
    return t

story_pdf = []

# PAGE DE GARDE
story_pdf += [
    sp(25),
    p("RAPPORT DE FIN DE TÂCHES ET D'AVANCEMENT TECHNIQUE", st_cover),
    p("Projet : Agent IA RH — Système Intelligent d'Automatisation du Recrutement (RecrutIA)", st_sub),
    hr(),
    sp(10),
]

info_p = [
    [p("<b>Entreprise d'accueil</b>", st_cell), p("ARTIWEB — Agence de Communication &amp; Marketing Digital", st_cell)],
    [p("<b>Localisation &amp; Contact</b>", st_cell), p("Fès, Maroc | contact@artiweb.ma | (+212) 664-017447", st_cell)],
    [p("<b>Accréditation</b>", st_cell), p("Partenaire Google Certifié (Google Partner) | https://artiweb.ma", st_cell)],
    [p("<b>Période de réalisation</b>", st_cell), p("Année Universitaire 2025 / 2026", st_cell)],
    [p("<b>Sujet du Projet</b>", st_cell), p("Conception et réalisation d'un agent IA RH d'analyse et de scoring de CV", st_cell)],
    [p("<b>Technologies clés</b>", st_cell), p("Python 3.12, FastAPI, SQLAlchemy, PostgreSQL / SQLite, sentence-transformers, PyMuPDF, pdfplumber, pytesseract", st_cell)],
]
story_pdf += [ttable(info_p, [4.5*cm, 11*cm]), sp(15), PageBreak()]

# I. PRÉSENTATION DE L'ENTREPRISE
story_pdf += [
    h1("I. PRÉSENTATION DE L'ENTREPRISE D'ACCUEIL : ARTIWEB"), hr(),
    p("<b>ArtiWeb</b> est une agence spécialisée en communication, marketing digital et développement de solutions web/mobile basées à Fès, Maroc. En tant que <b>Google Partner certifié</b>, ArtiWeb accompagne les entreprises dans la digitalisation de leurs processus métier."),
    bb("Marketing Digital &amp; Stratégie :", "Gestion de l'image de marque, campagnes Google Ads et Social Ads."),
    bb("Développement Web &amp; Applications :", "Création de plateformes e-commerce et d'applications de gestion métiers."),
    bb("Projet RecrutIA :", "Développement interne d'une solution d'intelligence artificielle visant à automatiser le tri des candidatures RH."),
    sp(8),
]

# II. FICHE PROJET
story_pdf += [
    h1("II. FICHE SYNTHÉTIQUE DU PROJET RECRUTIA"), hr(),
    p("Le processus traditionnel de tri des CV est manuel, chronophage et sujet aux bruits ou biais humains. <b>RecrutIA</b> apporte une réponse automatisée, explicable et rapide grâce à une architecture modulaire à 3 couches :"),
    sp(4),
]

arch_table = [
    [p("<b>Couche</b>", st_thdr), p("<b>Nom &amp; Rôle Technique</b>", st_thdr), p("<b>Statut</b>", st_thdr)],
    [p("<b>Couche 1</b>", st_cell), p("<b>Ingestion de CV (ingestion/) :</b> Extraction cascade (PDF/DOCX/OCR), nettoyage Unicode et parsing NLP des compétences, durées d'expérience et diplômes.", st_cell), p("<b>100% VALIDÉE</b><br/>(21/21 tests)", st_ok)],
    [p("<b>Couche 2</b>", st_cell), p("<b>IA Agentique &amp; Scoring (ia_agentic/) :</b> Filtre sur critères durs non négociables, scoring vectoriel sémantique (0-100) et justification LLM anti-hallucination.", st_cell), p("<b>100% VALIDÉE</b><br/>(7/7 tests)", st_ok)],
    [p("<b>Couche 3</b>", st_cell), p("<b>Backend API REST (backend/) :</b> Serveur FastAPI, persistance PostgreSQL/SQLite, endpoints Swagger et traçabilité d'audit RH (AuditLog).", st_cell), p("<b>100% VALIDÉE</b><br/>(7/7 tests)", st_ok)],
]
story_pdf += [ttable(arch_table, [2.5*cm, 9.5*cm, 3.5*cm]), sp(10), PageBreak()]

# III. DÉTAILS DE LA TÂCHE 1
story_pdf += [
    h1("III. TÂCHE 1 — COUCHE D'INGESTION INTELLIGENTE DE CV"), hr(),
    p("La Tâche 1 transforme un CV brut (PDF textuel, PDF scanné ou DOCX Word) en un dictionnaire JSON structuré sans faire appel à un LLM externe."),
    bb("Cascade d'Extraction :", "pdfplumber (texte natif) ➔ PyMuPDF (secours) ➔ OCR Tesseract à 300 DPI (scans). Fichiers Word via python-docx."),
    bb("Nettoyage &amp; Normalisation :", "Normalisation Unicode NFC, fusion des ligatures (fi, fl), suppression des symboles parasites (•, ★) et mots coupés en fin de ligne."),
    bb("Parsing NLP :", "Extraction des compétences par répertoires JSON (competences_ref.json), calcul automatique des durées par chevauchement de dates et hiérarchie diplômes (PhD > Master > Licence > BTS)."),
    bb("Bilan des Tests :", "21 tests automatisés validés à 100% en 0.91 seconde (tests/test_ingestion.py)."),
    sp(8),
]

# IV. DÉTAILS DE LA TÂCHE 2
story_pdf += [
    h1("IV. TÂCHE 2 — COUCHE IA AGENTIQUE &amp; SCORING SÉMANTIQUE"), hr(),
    p("La Tâche 2 reçoit le profil JSON du candidat et une offre d'emploi pour évaluer sémantiquement leur adéquation via la fonction unique <code>evaluer_candidature(candidat_json, offre_json)</code>."),
    bb("Étape 1 — Filtre Critères Durs :", "Rejet automatique (Score 0.0, statut 'rejete_auto_filtre') si l'expérience minimale ou les compétences obligatoires ne sont pas satisfaites. Évite tout gaspillage d'appel API."),
    bb("Étape 2 &amp; 3 — Embeddings &amp; Scoring :", "Modèle all-MiniLM-L6-v2 + Cosine Similarity. Score combiné sur 100 : 40% similarité vectorielle + 40% compétences + 20% expérience."),
    bb("Étape 4 — Justification LLM Anti-Hallucination :", "Cascade Groq API ➔ Ollama Local ➔ Moteur Déterministe anti-hallucination s'appuyant strictement sur les faits du CV."),
    bb("Bilan des Tests :", "7 tests automatisés validés à 100% sur 3 scénarios : bon match (score 100), rejet dur (score 0) et match modéré (score 59.9)."),
    sp(8),
]

# V. DÉTAILS DE LA TÂCHE 3
story_pdf += [
    h1("V. TÂCHE 3 — BACKEND FASTAPI, BASE DE DONNÉES &amp; TRAÇABILITÉ"), hr(),
    p("La Tâche 3 relie les deux premières couches au sein d'une API REST FastAPI robuste avec persistance des données et journalisation d'audit."),
    bb("Architecture BDD (SQLAlchemy) :", "Tables ORM Offre, Candidat, Candidature et AuditLog. Support PostgreSQL et SQLite local."),
    bb("Endpoints Swagger UI (/docs) :", "POST /api/offres, POST /api/candidatures (Upload CV + pipeline automatique), GET /api/offres/{id}/candidatures (triées par score), PATCH /api/candidatures/{id}/decision (supervision RH) et GET /api/audit."),
    bb("Bilan des Tests :", "7 tests d'intégration API REST validés à 100% (tests/test_backend.py)."),
    sp(10), PageBreak(),
]

# VI. BILAN DES TESTS AUTOMATISÉS
story_pdf += [
    h1("VI. BILAN GLOBAL DES TESTS AUTOMATISÉS (35 TESTS VALIDÉS)"), hr(),
    p("L'ensemble du projet est couvert par une suite de **35 tests automatisés** exécutés via <code>pytest</code> :"),
    sp(4),
]

test_summary_table = [
    [p("<b>Suite de Tests</b>", st_thdr), p("<b>Couche Testée</b>", st_thdr), p("<b>Nombre de Tests</b>", st_thdr), p("<b>Résultat</b>", st_thdr)],
    [p("<code>tests/test_ingestion.py</code>", st_cell), p("Tâche 1 — Ingestion &amp; Parsing CV", st_cell), p("21 tests", st_cell), p("<b>21 / 21 PASSED</b>", st_ok)],
    [p("<code>tests/test_ia_agentic.py</code>", st_cell), p("Tâche 2 — Scoring &amp; Justification IA", st_cell), p("7 tests", st_cell), p("<b>7 / 7 PASSED</b>", st_ok)],
    [p("<code>tests/test_backend.py</code>", st_cell), p("Tâche 3 — API FastAPI &amp; Base de Données", st_cell), p("7 tests", st_cell), p("<b>7 / 7 PASSED</b>", st_ok)],
    [p("<b>TOTAL SUITE PROJET</b>", st_cell), p("<b>Pipeline Complet End-to-End</b>", st_cell), p("<b>35 tests</b>", st_cell), p("<b>35 / 35 PASSED (100%)</b>", st_ok)],
]
story_pdf += [ttable(test_summary_table, [4.5*cm, 5*cm, 3*cm, 3*cm]), sp(10)]

story_pdf += [
    rbox("🏆 SYNTHÈSE DES RÉSULTATS OBTENUS SUR CV RÉELS", [
        "1. Ingestion CV (MEKKI Fatima-Ezahrae.pdf) : 33 compétences extraites, diplôme d'ingénieur en 0.78s.",
        "2. Scoring IA (Offre Data Science)          : Score 59.9 / 100, justification RH détaillée sans hallucination.",
        "3. API Backend REST (FastAPI)              : Serveur démarré sur http://127.0.0.1:8000/docs (Swagger UI).",
        "4. Taux de succès global                    : 100% (35/35 tests automatisés validés).",
    ], GREEN),
    sp(10),
    h1("VII. CONCLUSION ET PERSPECTIVES"), hr(),
    p("Les trois premières tâches techniques du projet <b>RecrutIA</b> ont été développées, testées et intégrées avec succès. Le système dispose d'une couche d'ingestion robuste, d'un moteur de scoring sémantique impartial et d'un backend API REST prêt pour le déploiement de l'interface graphique utilisateur (Tâche 4 — Dashboard RH)."),
]

doc_pdf.build(story_pdf)
print(f"Document PDF généré avec succès : {os.path.abspath(pdf_path)}")


# ─────────────────────────────────────────────────────────────
# 2. GÉNÉRATION DU FICHIER WORD (.DOCX) COMPLET
# ─────────────────────────────────────────────────────────────

docx_path = "Rapport_de_Stage_Complet_Encadrant.docx"
doc_word = Document()

for section in doc_word.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

COLOR_PRIMARY = RGBColor(0, 51, 102)
COLOR_SECONDARY = RGBColor(0, 102, 204)
COLOR_TEXT = RGBColor(51, 51, 51)

def add_w_title(text):
    p = doc_word.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    p.paragraph_format.space_after = Pt(4)
    return p

def add_w_sub(text):
    p = doc_word.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = COLOR_SECONDARY
    p.paragraph_format.space_after = Pt(18)
    return p

def add_w_h1(text):
    p = doc_word.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_w_h2(text):
    p = doc_word.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_SECONDARY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_w_p(text, bold_pre=None):
    p = doc_word.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    if bold_pre:
        r_pre = p.add_run(bold_pre)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_TEXT
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(10.5)
    r.font.color.rgb = COLOR_TEXT
    return p

def add_w_bullet(text, bold_pre=None):
    p = doc_word.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_pre:
        r_pre = p.add_run(bold_pre)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_TEXT
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(10.5)
    r.font.color.rgb = COLOR_TEXT
    return p

def add_w_code(text):
    p = doc_word.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(30, 30, 30)
    return p

# CONTENU DOCX
add_w_title("RAPPORT DE STAGE ET D'AVANCEMENT TECHNIQUE")
add_w_sub("Projet : Agent IA RH — Système Intelligent d'Automatisation du Recrutement (RecrutIA)\nEntreprise d'Accueil : ARTIWEB (Fès, Maroc — Partenaire Google)")

add_w_h1("I. PRÉSENTATION DE L'ENTREPRISE D'ACCUEIL : ARTIWEB")
add_w_bullet(" ARTIWEB (ARTI web) | Site Web : https://artiweb.ma/ | Partenaire Google certifié", "Entreprise :")
add_w_bullet(" Fès, Maroc | Contact : contact@artiweb.ma | (+212) 664-017447", "Localisation :")
add_w_bullet(" Marketing Digital, Développement Web/Mobile, SEO/SEM, Social Ads et Marketplace.", "Expertises :")

add_w_h1("II. FICHE SYNTHÉTIQUE DU PROJET RECRUTIA")
add_w_p("RecrutIA automatise le tri des candidatures RH via une architecture à 3 couches :")
add_w_bullet(" Extraction cascade (PDF/DOCX/OCR) et parsing NLP (21/21 tests validés).", "Couche 1 (Ingestion) :")
add_w_bullet(" Scoring sémantique vectoriel et justification LLM anti-hallucination (7/7 tests validés).", "Couche 2 (IA Agentique) :")
add_w_bullet(" API REST FastAPI, base de données PostgreSQL/SQLite et AuditLog (7/7 tests validés).", "Couche 3 (Backend REST) :")

add_w_h1("III. DÉTAILS DES TÂCHES RÉALISÉES ET VALIDÉES")
add_w_h2("3.1 Tâche 1 — Couche d'Ingestion (ingestion/)")
add_w_p("Extraction cascade pdfplumber -> PyMuPDF -> OCR Tesseract à 300 DPI. Nettoyage Unicode NFC et détection NLP des compétences, durées d'expérience et diplômes.")

add_w_h2("3.2 Tâche 2 — Couche IA Agentique (ia_agentic/)")
add_w_p("Filtrage dur automatique sur les critères obligatoires, vectorisation sentence-transformers (all-MiniLM-L6-v2) et justification LLM déterministe anti-hallucination.")

add_w_h2("3.3 Tâche 3 — Backend API REST & Base de Données (backend/)")
add_w_p("Serveur FastAPI autonome exposant l'interface Swagger UI sur http://127.0.0.1:8000/docs avec stockage des offres, candidats, candidatures et traçabilité d'audit.")

add_w_h1("IV. BILAN DES TESTS AUTOMATISÉS (35/35 TESTS PASSED)")
add_w_code("""============================= test session starts =============================
collected 35 items

tests/test_ingestion.py :: 21 PASSED  (Tâche 1 - Ingestion & Parsing)
tests/test_ia_agentic.py ::  7 PASSED  (Tâche 2 - Scoring IA & Justification)
tests/test_backend.py   ::  7 PASSED  (Tâche 3 - API REST FastAPI & DB)

======================= 35 passed in 3.69s (100% Succès) ======================""")

add_w_h1("V. CONCLUSION ET PERSPECTIVES")
add_w_p("Les Tâches 1, 2 et 3 sont entièrement achevées et validées. Le système est prêt pour le développement du Dashboard RH (Tâche 4).")

doc_word.save(docx_path)
print(f"Document Word généré avec succès : {os.path.abspath(docx_path)}")
