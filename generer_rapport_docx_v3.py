# -*- coding: utf-8 -*-
"""
generer_rapport_docx_v3.py
Génère le rapport Word (.docx) complet incluant Tâches 1, 2 et 3
"""
import sys, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

COLOR_PRIMARY = RGBColor(0, 51, 102)     # Navy
COLOR_SECONDARY = RGBColor(0, 102, 204)  # Royal Blue
COLOR_TEXT = RGBColor(51, 51, 51)

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    p.paragraph_format.space_after = Pt(4)
    return p

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = COLOR_SECONDARY
    p.paragraph_format.space_after = Pt(20)
    return p

def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COLOR_SECONDARY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_p(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_TEXT
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(10.5)
    r.font.color.rgb = COLOR_TEXT
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_TEXT
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(10.5)
    r.font.color.rgb = COLOR_TEXT
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(30, 30, 30)
    return p

# --- TITRE ---
add_title("RAPPORT DE STAGE ET D'AVANCEMENT TECHNIQUE")
add_subtitle("Projet Agent IA RH (RecrutIA) — Tâches 1 (Ingestion), 2 (IA Agentique) & 3 (Backend FastAPI & DB)\nEntreprise d'Accueil : ARTIWEB (Fès, Maroc — Partenaire Google)")

# --- SECTION 1 ---
add_h1("I. PRÉSENTATION ARTIWEB & PROJET RECRUTIA")
add_bullet(" ARTIWEB (ARTI web) | Site Web : https://artiweb.ma/ | Google Partner certifié", "Entreprise :")
add_bullet(" Tâche 1 (Ingestion) + Tâche 2 (IA Agentique & Scoring) + Tâche 3 (Backend API REST & DB)", "Architecture 3 Tâches :")

# --- SECTION 2 ---
add_h1("II. TÂCHE 1 : COUCHE D'INGESTION DE CV")
add_p("Extraction cascade PDF/DOCX/OCR 300 DPI, nettoyage Unicode NFC et parsing NLP des compétences, durées d'expérience et diplômes. (21/21 tests PASSED).")

# --- SECTION 3 ---
add_h1("III. TÂCHE 2 : COUCHE IA AGENTIQUE & SCORING SÉMANTIQUE")
add_p("Filtrage dur non négociable, scoring vectoriel (40% vectoriel + 40% compétences + 20% exp) et justification LLM anti-hallucination. (7/7 tests PASSED).")

# --- SECTION 4 ---
add_h1("IV. TÂCHE 3 : BACKEND FASTAPI, BASE DE DONNÉES & TRAÇABILITÉ")
add_h2("4.1 Description & Endpoints REST")
add_p("API REST FastAPI autonome exposant la documentation Swagger UI sur http://127.0.0.1:8000/docs et connectant la Tâche 1 et la Tâche 2 au SGBD PostgreSQL / SQLite.")
add_bullet(" Création d'une offre RH avec compétences requises et expérience min.", "POST /api/offres :")
add_bullet(" Upload du CV (.pdf/.docx) + offre_id. Déclenche le pipeline automatique Tâche 1 ➔ Tâche 2 ➔ DB.", "POST /api/candidatures :")
add_bullet(" Liste des candidatures d'une offre, TRIÉES PAR SCORE DÉCROISSANT.", "GET /api/offres/{id}/candidatures :")
add_bullet(" Enregistrer la décision finale (VALIDE, CORRIGE, REJETE) avec note et traçabilité d'audit.", "PATCH /api/candidatures/{id}/decision :")
add_bullet(" Journal d'audit complet de toutes les actions système et RH.", "GET /api/audit :")

add_h2("4.2 Validation des 35 Tests du Projet")
add_code("""============================= test session starts =============================
collected 35 items

tests/test_backend.py::TestBackendAPI::test_01_health_check PASSED
tests/test_backend.py::TestBackendAPI::test_02_creer_offre PASSED
tests/test_backend.py::TestBackendAPI::test_03_lister_offres PASSED
tests/test_backend.py::TestBackendAPI::test_04_soumettre_candidature_et_pipeline PASSED
tests/test_backend.py::TestBackendAPI::test_05_lister_candidatures_triees PASSED
tests/test_backend.py::TestBackendAPI::test_06_enregistrer_decision_rh PASSED
tests/test_backend.py::TestBackendAPI::test_07_consulter_audit_log PASSED
[... 21 tests Tâche 1 PASSED ...]
[... 7 tests Tâche 2 PASSED ...]

======================= 35 passed in 3.69s (100% Succès) ======================""")

doc.save("Rapport_de_Stage_Agent_IA_RH_ArtiWeb.docx")
print("DOCX v3 généré avec succès.")
