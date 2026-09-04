"""
Module : test_ingestion.py
Rôle   : Tests unitaires et d'intégration de la couche ingestion.
         Vérifie le pipeline complet sur 3 scénarios de CV différents.

Exécution :
    python -m pytest tests/test_ingestion.py -v
    # ou
    python tests/test_ingestion.py

Auteur : Agent IA RH — Couche Ingestion
"""

import sys
import os
import json
import tempfile
import unittest
from pathlib import Path

# Ajouter la racine du projet au chemin Python
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from ingestion.traiter_cv import traiter_cv
from ingestion.cleaner import nettoyer_texte
from ingestion.parser import (
    extraire_competences,
    extraire_experience_annees,
    extraire_formation,
)


# ════════════════════════════════════════════════════════════════
# Helpers — Création de CV de test synthétiques
# ════════════════════════════════════════════════════════════════

def creer_cv_pdf_bien_formate(chemin: str):
    """Crée un CV PDF bien formaté (texte natif, haute qualité)."""
    try:
        import fitz
    except ImportError:
        raise unittest.SkipTest("PyMuPDF non installé")

    doc = fitz.open()
    page = doc.new_page()

    contenu = """
CURRICULUM VITAE

Nom : Ahmed El Amrani
Email : ahmed.elamrani@email.com | Téléphone : +212 6 12 34 56 78

FORMATION
Master en Informatique — Université Mohammed V, Rabat (2020-2022)
Licence en Génie Logiciel — ENSIAS, Rabat (2017-2020)

EXPÉRIENCE PROFESSIONNELLE

Développeur Full Stack — TechMaroc Solutions (2022 - 2024)
  - Développement d'applications web avec React et Django
  - Intégration de bases de données PostgreSQL et MongoDB
  - Déploiement sur AWS avec Docker et CI/CD

Stagiaire Développeur Python — DataVision (2021)
  - Machine Learning avec Scikit-learn et Pandas
  - Analyse de données et visualisation avec Matplotlib

COMPÉTENCES TECHNIQUES
Langages    : Python, JavaScript, TypeScript, Java
Frameworks  : React, Django, Flask, FastAPI, Spring Boot
Bases       : PostgreSQL, MongoDB, MySQL, Redis
DevOps      : Docker, Kubernetes, GitHub Actions, AWS, Linux
IA/ML       : Machine Learning, NLP, TensorFlow, Pandas, NumPy
"""
    page.insert_text((50, 50), contenu, fontsize=11)
    doc.save(chemin)
    doc.close()


def creer_cv_pdf_mal_formate(chemin: str):
    """Crée un CV PDF mal formaté (caractères spéciaux, structure irrégulière)."""
    try:
        import fitz
    except ImportError:
        raise unittest.SkipTest("PyMuPDF non installé")

    doc = fitz.open()
    page = doc.new_page()

    # Texte avec caractères parasites et structure désordonnée
    contenu = """
C.V  |||  Fatima Zahra BENALI  |||  f.benali@mail.com
-----------------------------------------------------------

>>>> FORMATION <<<<
  . . . BTS Informatique --- OFPPT Casablanca (2019-2021)...
        + Baccalauréat Sciences Math -- Lycée Ibn Khaldoun (2019)

>>>> EXPERIENCE <<<<
Technicienne Informatique -- Groupe Maroc Telecom (Jan 2021 - Dec 2023)
  * Maintenance réseaux TCP/IP
  * Support technique Excel, Word, PowerPoint
  * Administration Linux (Ubuntu, CentOS)
  * Gestion des bases MySQL et SQL Server

Stage -- ITMNET (Juin 2020 - Août 2020)
  * PHP, HTML, CSS

>>>> OUTILS & COMPETENCES <<<<
PHP | HTML | CSS | MySQL | SQL Server | Linux | Excel | Git
Réseau : TCP/IP | Administration systèmes
"""
    page.insert_text((50, 50), contenu, fontsize=10)
    doc.save(chemin)
    doc.close()


def creer_cv_docx(chemin: str):
    """Crée un CV au format DOCX (Word)."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise unittest.SkipTest("python-docx non installé")

    doc = Document()

    doc.add_heading("CV — Youssef CHERKAOUI", 0)
    doc.add_paragraph("Email : y.cherkaoui@gmail.com | Casablanca, Maroc")

    doc.add_heading("Formation", 1)
    doc.add_paragraph("Ingénieur en Génie Informatique — École Mohammadia d'Ingénieurs (2015-2020)")

    doc.add_heading("Expérience Professionnelle", 1)
    doc.add_paragraph("Lead Developer — FinTech Maroc (2020 - 2024)\n"
                       "  - Architecture microservices avec Spring Boot et Docker\n"
                       "  - Développement mobile Kotlin (Android)\n"
                       "  - CI/CD avec Jenkins et GitHub Actions\n"
                       "  - Base de données PostgreSQL et Redis")
    doc.add_paragraph("Développeur Backend — StartupHub (2018-2020)\n"
                       "  - Python, Django, REST API\n"
                       "  - Deep Learning avec PyTorch\n"
                       "  - GCP, Kubernetes")

    doc.add_heading("Compétences Techniques", 1)
    tableau = doc.add_table(rows=3, cols=2)
    tableau.cell(0, 0).text = "Langages"
    tableau.cell(0, 1).text = "Python, Java, Kotlin, TypeScript"
    tableau.cell(1, 0).text = "Frameworks"
    tableau.cell(1, 1).text = "Spring Boot, Django, React"
    tableau.cell(2, 0).text = "Cloud & DevOps"
    tableau.cell(2, 1).text = "Docker, Kubernetes, GCP, Jenkins, GitHub Actions"

    doc.save(chemin)


# ════════════════════════════════════════════════════════════════
# Tests unitaires du Cleaner
# ════════════════════════════════════════════════════════════════

class TestCleaner(unittest.TestCase):

    def test_suppression_caracteres_parasites(self):
        texte = "Compétences : • Python • Java ★ Django"
        propre = nettoyer_texte(texte)
        self.assertNotIn("•", propre)
        self.assertNotIn("★", propre)
        self.assertIn("Python", propre)

    def test_normalisation_unicode(self):
        texte = "Exp\u00e9rience en d\u00e9veloppement"
        propre = nettoyer_texte(texte)
        self.assertIn("Expérience", propre)

    def test_ligatures_pdf(self):
        texte = "Pro\ufb01l : développeur"  # ﬁ → fi
        propre = nettoyer_texte(texte)
        self.assertIn("fi", propre)

    def test_espaces_multiples(self):
        texte = "Python   Django    React"
        propre = nettoyer_texte(texte)
        self.assertNotIn("   ", propre)

    def test_texte_vide(self):
        self.assertEqual(nettoyer_texte(""), "")
        self.assertEqual(nettoyer_texte("   "), "")


# ════════════════════════════════════════════════════════════════
# Tests unitaires du Parser
# ════════════════════════════════════════════════════════════════

class TestParser(unittest.TestCase):

    def test_extraction_competences_basique(self):
        texte = "Compétences : Python, Django, PostgreSQL, Docker, React"
        comps = extraire_competences(texte)
        self.assertIn("Python", comps)
        self.assertIn("Django", comps)
        self.assertIn("PostgreSQL", comps)

    def test_extraction_competences_insensible_casse(self):
        texte = "J'utilise PYTHON et javascript pour mes projets."
        comps = extraire_competences(texte)
        self.assertIn("Python", comps)
        self.assertIn("JavaScript", comps)

    def test_experience_mention_explicite(self):
        texte = "Développeur avec 5 ans d'expérience en Python."
        exp = extraire_experience_annees(texte)
        self.assertEqual(exp, 5)

    def test_experience_calcul_dates(self):
        texte = "Développeur Senior — TechCorp (2019 - 2023)\nDéveloppeur Junior (2017 - 2019)"
        exp = extraire_experience_annees(texte)
        # 2019-2023 = 4 ans, 2017-2019 = 2 ans → total 6 ans
        self.assertGreaterEqual(exp, 4)

    def test_experience_zero_si_non_trouve(self):
        texte = "Développeur passionné par la technologie."
        exp = extraire_experience_annees(texte)
        self.assertEqual(exp, 0)

    def test_formation_master(self):
        texte = "Master en Informatique — Université de Rabat (2020)"
        form = extraire_formation(texte)
        self.assertEqual(form, "Master / Ingénieur")

    def test_formation_ingenieur(self):
        texte = "Diplôme d'Ingénieur — École Mohammadia (2022)"
        form = extraire_formation(texte)
        self.assertEqual(form, "Master / Ingénieur")

    def test_formation_bts(self):
        texte = "BTS Informatique de Gestion — OFPPT (2021)"
        form = extraire_formation(texte)
        self.assertEqual(form, "BTS / DUT")

    def test_formation_non_specifiee(self):
        texte = "Développeur autodidacte avec de nombreux projets personnels."
        form = extraire_formation(texte)
        self.assertEqual(form, "Non spécifié")

    def test_formation_priorite_plus_haut_niveau(self):
        # Doctorat > Master : doit retourner Doctorat
        texte = "PhD en Machine Learning — Master en Informatique — Baccalauréat"
        form = extraire_formation(texte)
        self.assertEqual(form, "Doctorat / PhD")


# ════════════════════════════════════════════════════════════════
# Tests d'intégration — Pipeline complet traiter_cv()
# ════════════════════════════════════════════════════════════════

class TestIntegrationTraiterCV(unittest.TestCase):

    def setUp(self):
        """Crée un répertoire temporaire pour les fichiers de test."""
        self.tmpdir = tempfile.mkdtemp()

    def _verifier_structure_resultat(self, resultat: dict):
        """Vérifie que le dictionnaire retourné a bien la structure attendue."""
        self.assertIsInstance(resultat, dict)
        self.assertIn("competences", resultat)
        self.assertIn("experience_annees", resultat)
        self.assertIn("formation", resultat)
        self.assertIn("texte_brut", resultat)
        self.assertIn("statut", resultat)
        self.assertIsInstance(resultat["competences"], list)
        self.assertIsInstance(resultat["experience_annees"], int)
        self.assertIsInstance(resultat["formation"], str)
        self.assertIsInstance(resultat["texte_brut"], str)

    def test_cv_pdf_bien_formate(self):
        """Test 1 — CV PDF bien formaté (texte natif)"""
        chemin = os.path.join(self.tmpdir, "cv_bien_formate.pdf")
        creer_cv_pdf_bien_formate(chemin)

        print("\n" + "=" * 60)
        print("TEST 1 : CV PDF bien formaté")
        print("=" * 60)

        resultat = traiter_cv(chemin)
        self._verifier_structure_resultat(resultat)

        print(f"  Statut          : {resultat['statut']}")
        print(f"  Formation       : {resultat['formation']}")
        print(f"  Expérience      : {resultat['experience_annees']} an(s)")
        print(f"  Compétences     : {resultat['competences']}")
        print(f"  Durée           : {resultat['duree_traitement']}s")

        self.assertEqual(resultat["statut"], "succès")
        self.assertGreater(len(resultat["texte_brut"]), 50)
        self.assertGreater(len(resultat["competences"]), 0)

        # Vérifications spécifiques au contenu du CV
        self.assertIn("Python", resultat["competences"])
        self.assertIn("React", resultat["competences"])
        self.assertIn("Master / Ingénieur", resultat["formation"])

    def test_cv_pdf_mal_formate(self):
        """Test 2 — CV PDF mal formaté (symboles, structure désordonnée)"""
        chemin = os.path.join(self.tmpdir, "cv_mal_formate.pdf")
        creer_cv_pdf_mal_formate(chemin)

        print("\n" + "=" * 60)
        print("TEST 2 : CV PDF mal formaté")
        print("=" * 60)

        resultat = traiter_cv(chemin)
        self._verifier_structure_resultat(resultat)

        print(f"  Statut          : {resultat['statut']}")
        print(f"  Formation       : {resultat['formation']}")
        print(f"  Expérience      : {resultat['experience_annees']} an(s)")
        print(f"  Compétences     : {resultat['competences']}")
        print(f"  Durée           : {resultat['duree_traitement']}s")

        self.assertEqual(resultat["statut"], "succès")
        self.assertIn("PHP", resultat["competences"])
        self.assertIn("MySQL", resultat["competences"])
        self.assertIn("Linux", resultat["competences"])
        self.assertEqual(resultat["formation"], "BTS / DUT")

    def test_cv_docx(self):
        """Test 3 — CV DOCX (format Word avec tableau)"""
        chemin = os.path.join(self.tmpdir, "cv_word.docx")
        creer_cv_docx(chemin)

        print("\n" + "=" * 60)
        print("TEST 3 : CV DOCX (Word)")
        print("=" * 60)

        resultat = traiter_cv(chemin)
        self._verifier_structure_resultat(resultat)

        print(f"  Statut          : {resultat['statut']}")
        print(f"  Formation       : {resultat['formation']}")
        print(f"  Expérience      : {resultat['experience_annees']} an(s)")
        print(f"  Compétences     : {resultat['competences']}")
        print(f"  Durée           : {resultat['duree_traitement']}s")

        self.assertEqual(resultat["statut"], "succès")
        self.assertIn("Python", resultat["competences"])
        self.assertIn("Docker", resultat["competences"])
        self.assertIn("Master / Ingénieur", resultat["formation"])

    def test_fichier_inexistant(self):
        """Test gestion d'erreur — fichier introuvable"""
        resultat = traiter_cv("/chemin/inexistant/cv.pdf")
        self.assertEqual(resultat["statut"], "erreur")
        self.assertIn("introuvable", resultat["message"].lower())

    def test_format_non_supporte(self):
        """Test gestion d'erreur — format non supporté"""
        chemin = os.path.join(self.tmpdir, "cv.xlsx")
        with open(chemin, "w") as f:
            f.write("dummy")
        resultat = traiter_cv(chemin)
        self.assertEqual(resultat["statut"], "erreur")

    def test_affichage_json_final(self):
        """Affiche le résultat JSON final pour visualisation."""
        chemin = os.path.join(self.tmpdir, "cv_demo.pdf")
        creer_cv_pdf_bien_formate(chemin)
        resultat = traiter_cv(chemin)

        # Retirer le texte brut pour l'affichage (trop long)
        affichage = {k: v for k, v in resultat.items() if k != "texte_brut"}
        affichage["texte_brut"] = f"[{len(resultat['texte_brut'])} caractères]"

        print("\n" + "=" * 60)
        print("RÉSULTAT JSON FINAL")
        print("=" * 60)
        print(json.dumps(affichage, ensure_ascii=False, indent=2))


# ════════════════════════════════════════════════════════════════
# Point d'entrée
# ════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("\n" + "=" * 60)
    print("   TESTS -- Couche Ingestion | Agent IA RH")
    print("=" * 60)
    unittest.main(verbosity=2)
