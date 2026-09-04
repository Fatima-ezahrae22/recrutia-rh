"""
Module : extractor.py
Rôle   : Extraction du texte brut depuis un fichier CV (PDF ou DOCX).
         Mécanisme de secours OCR activé si le texte natif est absent ou insuffisant.

Auteur : Agent IA RH — Couche Ingestion
"""

import os
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────
# 1. Extraction PDF (texte natif)
# ─────────────────────────────────────────

def _extraire_pdf_pdfplumber(chemin: str) -> str:
    """Extrait le texte d'un PDF avec pdfplumber (couche texte native)."""
    try:
        import pdfplumber
        texte_pages = []
        with pdfplumber.open(chemin) as pdf:
            for page in pdf.pages:
                texte = page.extract_text()
                if texte:
                    texte_pages.append(texte)
        return "\n".join(texte_pages)
    except Exception as e:
        logger.warning(f"[pdfplumber] Échec sur '{chemin}': {e}")
        return ""


def _extraire_pdf_pymupdf(chemin: str) -> str:
    """Extrait le texte d'un PDF avec PyMuPDF (fitz) — secours niveau 1."""
    try:
        import fitz  # PyMuPDF
        texte_pages = []
        with fitz.open(chemin) as doc:
            for page in doc:
                texte_pages.append(page.get_text())
        return "\n".join(texte_pages)
    except Exception as e:
        logger.warning(f"[PyMuPDF] Échec sur '{chemin}': {e}")
        return ""


# ─────────────────────────────────────────
# 2. OCR — Secours pour PDF scannés
# ─────────────────────────────────────────

def _extraire_pdf_ocr(chemin: str) -> str:
    """
    Convertit chaque page PDF en image puis applique Tesseract OCR.
    Utilisé uniquement si les méthodes natives échouent ou retournent < 50 caractères.
    """
    try:
        import fitz
        import pytesseract
        from PIL import Image
        import io

        # Tesseract doit être installé sur le système Windows
        # Chemin par défaut : C:/Program Files/Tesseract-OCR/tesseract.exe
        tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path

        texte_pages = []
        with fitz.open(chemin) as doc:
            for i, page in enumerate(doc):
                # Rendu haute résolution (300 dpi) pour meilleure précision OCR
                mat = fitz.Matrix(300 / 72, 300 / 72)
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_bytes))

                # OCR avec langues française et anglaise
                config = "--psm 6 --oem 3"
                texte = pytesseract.image_to_string(img, lang="fra+eng", config=config)
                texte_pages.append(texte)
                logger.info(f"[OCR] Page {i+1}/{len(doc)} traitée.")

        return "\n".join(texte_pages)

    except ImportError as e:
        logger.error(f"[OCR] Dépendance manquante : {e}. Installez pytesseract + Tesseract.")
        return ""
    except Exception as e:
        logger.error(f"[OCR] Erreur inattendue : {e}")
        return ""


# ─────────────────────────────────────────
# 3. Extraction DOCX
# ─────────────────────────────────────────

def _extraire_docx(chemin: str) -> str:
    """Extrait le texte d'un fichier DOCX en parcourant tous les paragraphes."""
    try:
        from docx import Document
        doc = Document(chemin)
        paragraphes = [p.text for p in doc.paragraphs if p.text.strip()]
        # Inclure également les tableaux (souvent utilisés dans les CV)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphes.append(cell.text.strip())
        return "\n".join(paragraphes)
    except Exception as e:
        logger.error(f"[DOCX] Erreur lors de l'extraction de '{chemin}': {e}")
        return ""


# ─────────────────────────────────────────
# 4. Dispatcher principal
# ─────────────────────────────────────────

# Seuil minimal de caractères pour considérer une extraction réussie
SEUIL_TEXTE_MINIMAL = 80


def extraire_texte(chemin_fichier: str) -> str:
    """
    Fonction principale d'extraction.
    Stratégie en cascade :
      PDF  → pdfplumber  → PyMuPDF  → OCR Tesseract
      DOCX → python-docx

    Args:
        chemin_fichier (str): Chemin absolu ou relatif vers le CV.

    Returns:
        str: Texte brut extrait, ou chaîne vide si échec total.

    Raises:
        FileNotFoundError : Si le fichier n'existe pas.
        ValueError        : Si le format de fichier n'est pas supporté.
    """
    chemin = Path(chemin_fichier).resolve()

    if not chemin.exists():
        raise FileNotFoundError(f"Fichier introuvable : {chemin}")

    extension = chemin.suffix.lower()
    logger.info(f"[Extractor] Traitement de '{chemin.name}' (format: {extension})")

    # ── Traitement DOCX ──
    if extension == ".docx":
        texte = _extraire_docx(str(chemin))
        if len(texte.strip()) < SEUIL_TEXTE_MINIMAL:
            logger.warning("[DOCX] Texte insuffisant extrait.")
        return texte

    # ── Traitement PDF ──
    elif extension == ".pdf":
        # Tentative 1 : pdfplumber
        texte = _extraire_pdf_pdfplumber(str(chemin))
        if len(texte.strip()) >= SEUIL_TEXTE_MINIMAL:
            logger.info("[Extractor] Texte extrait via pdfplumber ✓")
            return texte

        logger.info("[Extractor] pdfplumber insuffisant → tentative PyMuPDF")

        # Tentative 2 : PyMuPDF
        texte = _extraire_pdf_pymupdf(str(chemin))
        if len(texte.strip()) >= SEUIL_TEXTE_MINIMAL:
            logger.info("[Extractor] Texte extrait via PyMuPDF ✓")
            return texte

        logger.info("[Extractor] PyMuPDF insuffisant → déclenchement OCR")

        # Tentative 3 : OCR (PDF scanné)
        texte = _extraire_pdf_ocr(str(chemin))
        if len(texte.strip()) >= SEUIL_TEXTE_MINIMAL:
            logger.info("[Extractor] Texte extrait via OCR Tesseract ✓")
        else:
            logger.error("[Extractor] Échec total : texte non récupérable.")
        return texte

    # ── Format non supporté ──
    else:
        raise ValueError(
            f"Format non supporté : '{extension}'. "
            f"Formats acceptés : .pdf, .docx"
        )
