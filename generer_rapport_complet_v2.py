# -*- coding: utf-8 -*-
"""
generer_rapport_complet_v2.py
Génère le rapport mis à jour incluant TÂCHE 1 (Ingestion) ET TÂCHE 2 (IA Agentique & Scoring)
Formats : DOCX + PDF
"""
import sys, os, json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# ── 1. Création du DOCX ──
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

COLOR_PRIMARY = RGBColor(0, 51, 102)     # Navy
COLOR_SECONDARY = RGBColor(0, 102, 204)  # Royal Blue
COLOR_TEXT = RGBColor(51, 51, 51)

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    p.paragraph_format.space_after = Pt(4)
    return p

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = COLOR_SECONDARY
    p.paragraph_format.space_after = Pt(20)
    return p

def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COLOR_SECONDARY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_p(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_TEXT
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(10.5)
    r.font.color.rgb = COLOR_TEXT
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_TEXT
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(10.5)
    r.font.color.rgb = COLOR_TEXT
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(30, 30, 30)
    return p

# --- TITRE ---
add_title("RAPPORT DE STAGE ET D'AVANCEMENT TECHNIQUE")
add_subtitle("Projet Agent IA RH (RecrutIA) — Modules 1 (Ingestion) & 2 (IA Agentique & Scoring)\nEntreprise d'Accueil : ARTIWEB (Fès, Maroc — Partenaire Google)")

# --- SECTION 1 ---
add_h1("I. PRÉSENTATION DE L'ENTREPRISE D'ACCUEIL : ARTIWEB")
add_bullet(" ARTIWEB (ARTI web) | Site Web : https://artiweb.ma/", "Nom de l'entreprise :")
add_bullet(" Fès, Maroc | Contacts : contact@artiweb.ma | (+212) 664-017447", "Siège & Contact :")
add_bullet(" Partenaire Google certifié (Google Partner)", "Accréditations :")
add_bullet(" Stratégie digitale, Développement Web & Mobile, SEO/SEM, Social Ads, Marketplace & Studio Marketing.", "Domaines d'expertise :")
add_p("ArtiWeb accompagne la transformation numérique des entreprises et développe en interne le projet innovant Agent IA RH (RecrutIA).")

# --- SECTION 2 ---
add_h1("II. FICHE PROJET : AGENT IA RH (RECRUTIA)")
add_p("RecrutIA est une solution d'automatisation intelligente du processus de recrutement visant à fluidifier le tri des candidatures et à assister les recruteurs RH.")
add_bullet(" Transformer les CV bruts (PDF/DOCX/Scan) en dictionnaire JSON structuré. (RÉALISÉE)", "Couche 1 - Ingestion :")
add_bullet(" Évaluer sémantiquement l'adéquation d'un candidat avec une offre, calculer un score sur 100 et générer une justification argumentée sans hallucination. (RÉALISÉE)", "Couche 2 - IA Agentique & Scoring :")
add_bullet(" Interface RH permettant de valider, corriger ou rejeter les préconisations IA (Human-in-the-loop). (À venir)", "Couche 3 - Supervision Humaine :")

# --- SECTION 3 ---
add_h1("III. TÂCHE 1 : COUCHE D'INGESTION DE CV (RÉALISÉE & VALIDÉE)")
add_p("Module Python autonome (ingestion/) convertissant tout CV brut en JSON structuré :")
add_bullet(" Cascade pdfplumber -> PyMuPDF -> Fallback OCR Tesseract (300 DPI) si scan < 80 chars.", "1. Extraction :")
add_bullet(" Normalisation Unicode NFC, suppression du bruit visuel (•, ★), fusion des mots coupés.", "2. Nettoyage :")
add_bullet(" Matching avec competences_ref.json, calcul des durées d'expérience et hiérarchie des diplômes.", "3. Parsing NLP :")
add_bullet(" 21/21 tests validés avec succès (100% de réussite).", "Validation :")

# --- SECTION 4 ---
add_h1("IV. TÂCHE 2 : COUCHE IA AGENTIQUE & SCORING SÉMANTIQUE (RÉALISÉE)")

add_h2("4.1 Objectif et Enjeux")
add_p("Développer un module Python autonome (ia_agentic/) capable de prendre en entrée le JSON du candidat et une offre d'emploi, puis de renvoyer la signature unique exigée :")
add_code("evaluer_candidature(candidat_json, offre_json) -> {\n  \"score\": float,\n  \"justification\": \"...\",\n  \"statut\": \"score\" | \"rejete_auto_filtre\"\n}")

add_h2("4.2 Architecture des Fichiers du Module (ia_agentic/)")
add_code("""agent_ai/
├── ia_agentic/
│   ├── __init__.py           # Export officiel de evaluer_candidature()
│   ├── hard_filter.py        # Étape 1 : Filtre sur critères durs (expérience min, compétences obligatoires)
│   ├── embeddings_scorer.py  # Étapes 2 & 3 : Embeddings multilingues + cosine similarity + score combiné
│   ├── llm_justifier.py      # Étape 4 : Justification LLM (Groq API / Ollama / Fallback déterministe)
│   └── evaluer_candidature.py# Étape 5 : Orchestrateur central
├── data/
│   └── offres_exemples.json  # 3 Offres d'emploi types (Backend, Data Science, DevOps)
├── tests/
│   ├── test_ingestion.py     # 21 tests Couche 1
│   └── test_ia_agentic.py    # 7 tests Couche 2 (Bon match, Mauvais match, Rejet dur)
└── run_ia_eval.py            # Démonstration globale de la Couche 2""")

add_h2("4.3 Les 5 Sous-Étapes de la Couche IA Agentique")
add_bullet(" Avant tout calcul lourd, vérification des critères non négociables (expérience minimale et compétences obligatoires). Si échec -> statut 'rejete_auto_filtre', score 0.0 et économie des appels API.", "Étape 1 - Filtre sur critères durs (hard_filter.py) :")
add_bullet(" Conversion du profil et de l'offre en vecteurs numériques comparables (sentence-transformers / cosine similarity).", "Étape 2 - Vectorisation Embeddings (embeddings_scorer.py) :")
add_bullet(" Combinaison pondérée : 40% similarité vectorielle, 40% couverture des compétences, 20% expérience.", "Étape 3 - Calcul du score combiné (0 à 100) :")
add_bullet(" Génération d'une explication RH claire via Groq API, Ollama ou le moteur déterministe anti-hallucination qui ne s'appuie que sur des faits réels du CV.", "Étape 4 - Génération de la justification LLM (llm_justifier.py) :")
add_bullet(" Enchaînement propre des étapes dans evaluer_candidature() avec mesure du temps.", "Étape 5 - Orchestration (evaluer_candidature.py) :")

# --- SECTION 5 ---
add_h1("V. RÉSULTATS DES TESTS AUTOMATISÉS ET VALIDATION EN CONDITION RÉELLE")

add_h2("5.1 Résultats de la Suite Globale de Tests (28 Tests)")
add_p("Exécution de la suite complète pytest (Tâche 1 + Tâche 2) :")
add_code("""============================= test session starts =============================
collected 28 items

tests/test_ia_agentic.py::TestHardFilter::test_candidat_rejete_competences_obligatoires_manquantes PASSED
tests/test_ia_agentic.py::TestHardFilter::test_candidat_rejete_experience_insuffisante PASSED
tests/test_ia_agentic.py::TestHardFilter::test_candidat_valide_criteres_durs PASSED
tests/test_ia_agentic.py::TestEmbeddingsScorer::test_score_bon_match_superieur_score_moyen PASSED
tests/test_ia_agentic.py::TestIntegrationEvaluerCandidature::test_scenario_1_bon_match PASSED
tests/test_ia_agentic.py::TestIntegrationEvaluerCandidature::test_scenario_2_rejet_automatique_filtre_dur PASSED
tests/test_ia_agentic.py::TestIntegrationEvaluerCandidature::test_scenario_3_cas_limite_match_modere PASSED
[... 21 tests de la Tâche 1 Ingestion PASSED ...]

============================= 28 passed in 1.05s ==============================""")

add_h2("5.2 Validation sur le CV Réel de Fatima-Ezahrae MEKKI")
add_p("Évaluation du CV ingéré (MEKKI_FATIMA-EZAHRAE_AI_Backend.pdf) face à 3 offres types :")

add_bullet(" STATUT: REJETE_AUTO_FILTRE | SCORE: 0.0/100 | Motif : Expérience insuffisante (2 ans vs 3 ans exigés).", "1. Offre Backend Python Senior (3 ans exp min) :")
add_bullet(" STATUT: SCORE | SCORE: 59.9/100 | Validation de la formation Ingénieure, de l'expérience et de 5 compétences clés (Python, TensorFlow, Scikit-learn, NLP, Computer Vision). Aucune hallucination.", "2. Offre Data Science & AI (1 an exp min) :")
add_bullet(" STATUT: REJETE_AUTO_FILTRE | SCORE: 0.0/100 | Motif : Expérience insuffisante (2 ans vs 5 ans) et compétences obligatoires manquantes (Kubernetes, AWS).", "3. Offre DevOps Senior (5 ans exp min) :")

# --- SECTION 6 ---
add_h1("VI. CONCLUSION ET PROCHAINES ÉTAPES")
add_p("Les Tâches 1 (Ingestion) et 2 (IA Agentique & Scoring) sont 100% terminées, testées (28 tests) et validées sur des CV réels. Le système est prêt pour le développement de la Tâche 3 (Backend & API Orchestratrice).")

doc.save("Rapport_de_Stage_Agent_IA_RH_ArtiWeb.docx")
print("Fichier DOCX v2 généré avec succès.")
